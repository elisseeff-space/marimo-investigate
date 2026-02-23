"""
Document reader with markdown conversion for .docx and .pdf files.

Converts binary document formats to markdown-formatted text for LLM processing.
Uses heading-based structure to enable intelligent chunking.
"""

from __future__ import annotations

from pathlib import Path

__all__ = [
    "is_docx_available",
    "is_pdf_available",
    "read_document",
    "read_docx",
    "read_pdf",
]


def is_docx_available() -> bool:
    """Check if python-docx library is installed."""
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


def is_pdf_available() -> bool:
    """Check if pypdf library is installed."""
    try:
        import pypdf  # noqa: F401

        return True
    except ImportError:
        return False


def read_document(file_path: Path) -> str | None:
    """
    Extract text from a document file as markdown-formatted text.

    Routes to appropriate reader based on file extension.
    Returns None if the required library is not installed or reading fails.

    Args:
        file_path: Path to the document file

    Returns:
        Markdown-formatted text content, or None if reading fails
    """
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        return read_docx(file_path)
    elif suffix == ".pdf":
        return read_pdf(file_path)
    else:
        return None


def read_docx(file_path: Path) -> str | None:
    """
    Extract text from a .docx file as markdown-formatted text.

    Preserves document structure:
    - Heading styles → markdown headings (#, ##, ###, etc.)
    - Bold/Italic → **bold**, *italic*
    - Bullet/Numbered lists → - item, 1. item
    - Tables → simple markdown table format
    - Paragraphs → text with blank lines

    Args:
        file_path: Path to the .docx file

    Returns:
        Markdown-formatted text, or None if reading fails
    """
    if not is_docx_available():
        return None

    try:
        import docx

        doc = docx.Document(str(file_path))
        lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check paragraph style for headings
            style_name = para.style.name if para.style else ""

            if style_name.startswith("Heading"):
                # Extract heading level (Heading 1, Heading 2, etc.)
                try:
                    level = int(style_name.split()[-1])
                    prefix = "#" * level
                    lines.append(f"{prefix} {text}")
                except (ValueError, IndexError):
                    # Fallback for non-standard heading styles
                    lines.append(f"## {text}")
            elif style_name == "Title":
                lines.append(f"# {text}")
            elif style_name.startswith("List"):
                # Handle list items
                lines.append(f"- {text}")
            else:
                # Regular paragraph - apply inline formatting
                formatted = _format_paragraph_runs(para)
                lines.append(formatted)

        # Process tables
        for table in doc.tables:
            table_md = _table_to_markdown(table)
            if table_md:
                lines.append("")
                lines.append(table_md)

        return "\n\n".join(lines)

    except Exception:
        return None


def _format_paragraph_runs(para) -> str:
    """
    Format a paragraph with inline styles (bold, italic).

    Args:
        para: python-docx Paragraph object

    Returns:
        Formatted text with markdown inline styles
    """
    parts: list[str] = []

    for run in para.runs:
        text = run.text
        if not text:
            continue

        # Apply formatting
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"

        parts.append(text)

    return "".join(parts)


def _table_to_markdown(table) -> str:
    """
    Convert a docx table to markdown format.

    Args:
        table: python-docx Table object

    Returns:
        Markdown-formatted table
    """
    rows: list[str] = []

    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        row_str = "| " + " | ".join(cells) + " |"
        rows.append(row_str)

        # Add header separator after first row
        if i == 0:
            separator = "| " + " | ".join(["---"] * len(cells)) + " |"
            rows.append(separator)

    return "\n".join(rows)


def read_pdf(file_path: Path) -> str | None:
    """
    Extract text from a .pdf file as markdown-formatted text.

    Structure:
    - Each page becomes a section: ## Page N
    - Text is extracted preserving paragraph breaks where possible
    - Tables are converted to plain text (limited structure preservation)

    Args:
        file_path: Path to the .pdf file

    Returns:
        Markdown-formatted text, or None if reading fails
    """
    if not is_pdf_available():
        return None

    try:
        import pypdf

        reader = pypdf.PdfReader(str(file_path))
        sections: list[str] = []

        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                # Clean up the text
                text = text.strip()
                # Add page header for chunking
                sections.append(f"## Page {i}\n\n{text}")

        return "\n\n".join(sections)

    except Exception:
        return None
